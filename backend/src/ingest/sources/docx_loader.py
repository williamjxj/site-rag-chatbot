"""Word document loader (.docx)."""

from pathlib import Path
from docx import Document


def load_docx(path: Path) -> dict[str, str]:
    """
    Load and extract text from Word .docx file.

    Args:
        path: Path to .docx file

    Returns:
        Dictionary with 'uri', 'title', and 'text' keys
    """
    doc = Document(str(path))
    parts = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))
    
    text = "\n".join(parts).strip()
    return {"uri": str(path), "title": path.stem, "text": text}

